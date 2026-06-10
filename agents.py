#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
מערכת סוכני AI לעסק אונליין בשוק האמריקאי
==========================================
"אימפריית סוכנים" קטנה ועובדת. הצוות:

  🧑‍💼 מנכ"ל     - מפרק מטרה עסקית למשימות מחקר ומסכם דוח החלטות
  🔍 חוקרים     - מחפשים בשוק האמריקאי האמיתי (מחירים, מתחרים, ביקוש)
  📦 ספקים      - מחפשים ב-Alibaba/AliExpress ומחשבים עלות נחיתה אמיתית
  💰 כלכלן      - מחשב עלויות, תמחור ורווחיות למכירה אונליין
  ✍️  שיווק      - כותב תוכן שיווקי באנגלית (לשוק האמריקאי) + הסבר בעברית

הפעלה דרך שורת הפקודה:  python3 agents.py
הפעלה דרך דפדפן:        python3 app.py
ראה README.md להוראות צעד-צעד.
"""

import os
import sys
import json
import datetime

try:
    import anthropic
except ImportError:
    print("חסרה ספריית anthropic. הרץ קודם:  pip install -r requirements.txt")
    sys.exit(1)


# ---------------------------------------------------------------------------
# הגדרות
# ---------------------------------------------------------------------------

# בחירת מודל דרך משתנה סביבה AGENT_MODEL (בלי לגעת בקוד):
#   "claude-sonnet-4-6"  -> ברירת מחדל. האיזון הכי טוב בין איכות למחיר
#   "claude-opus-4-8"    -> הכי חכם, יקר יותר. להחלטות גדולות
#   "claude-haiku-4-5"   -> הכי זול ומהיר, פחות מעמיק
MODEL = os.environ.get("AGENT_MODEL", "claude-sonnet-4-6")

WEB_SEARCH_TOOL = {"type": "web_search_20260209", "name": "web_search"}

# כמה שאלות מחקר המנכ"ל יפרק את המטרה (אפשר לשנות)
NUM_RESEARCH_TASKS = 4


def get_client() -> "anthropic.Anthropic":
    """יוצר חיבור ל-Claude. דורש משתנה סביבה ANTHROPIC_API_KEY."""
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print(
            "\n[שגיאה] לא נמצא מפתח API.\n"
            "הגדר אותו כך (החלף את xxxxx במפתח שלך מ-console.anthropic.com):\n"
            "    Mac/Linux:   export ANTHROPIC_API_KEY=sk-ant-xxxxx\n"
            "    Windows:     set ANTHROPIC_API_KEY=sk-ant-xxxxx\n"
        )
        sys.exit(1)
    return anthropic.Anthropic()


# ---------------------------------------------------------------------------
# מנוע בסיסי: הרצת סוכן בודד (עם או בלי גישה לאינטרנט)
# ---------------------------------------------------------------------------

def run_agent(client, system_prompt, user_prompt, use_web=False, max_tokens=16000):
    """
    מריץ סוכן Claude יחיד ומחזיר את הטקסט שהוא הפיק.
    אם use_web=True - הסוכן יכול לחפש באינטרנט בזמן אמת.
    """
    tools = [WEB_SEARCH_TOOL] if use_web else []
    messages = [{"role": "user", "content": user_prompt}]

    # לולאה שמטפלת ב-pause_turn (קורה כשהסוכן עושה הרבה חיפושים ברצף)
    for _ in range(10):
        resp = client.messages.create(
            model=MODEL,
            max_tokens=max_tokens,
            thinking={"type": "adaptive"},
            output_config={"effort": "high"},
            system=system_prompt,
            messages=messages,
            tools=tools,
        )
        if resp.stop_reason == "pause_turn":
            messages.append({"role": "assistant", "content": resp.content})
            continue
        break

    return "".join(block.text for block in resp.content if block.type == "text").strip()


# ---------------------------------------------------------------------------
# הסוכנים
# ---------------------------------------------------------------------------

def ceo_plan(client, business_goal):
    """המנכ"ל: מפרק מטרה עסקית ל-NUM_RESEARCH_TASKS שאלות מחקר."""
    system = (
        "אתה מנכ\"ל מנוסה של חברת ייעוץ לעסקים אונליין שמתמחה בשוק האמריקאי. "
        "הלקוח שלך מוכר אונליין בלבד ואינו דובר אנגלית, לכן התמקד בהזדמנויות "
        "שלא דורשות שיחה עם לקוחות (מסחר אלקטרוני, דרופשיפינג, מוצרים דיגיטליים). "
        "המשימה שלך: לפרק מטרה עסקית למשימות מחקר ממוקדות שאפשר לבדוק עם חיפוש "
        "באינטרנט — מחירים, מתחרים, גודל שוק, ביקוש, ועלויות כניסה נמוכות."
    )
    user = (
        f"המטרה העסקית: {business_goal}\n\n"
        f"פרק אותה ל-{NUM_RESEARCH_TASKS} שאלות מחקר עצמאיות. "
        "החזר אך ורק JSON תקין במבנה הבא, בלי טקסט נוסף:\n"
        '{"tasks": ["שאלה 1", "שאלה 2", ...]}'
    )
    raw = run_agent(client, system, user, use_web=False, max_tokens=4000)

    try:
        start = raw.index("{")
        end = raw.rindex("}") + 1
        tasks = json.loads(raw[start:end]).get("tasks", [])
        if tasks:
            return tasks[:NUM_RESEARCH_TASKS]
    except (ValueError, json.JSONDecodeError):
        pass

    return [
        f"חקור את הביקוש וגודל השוק האמריקאי עבור: {business_goal}",
        f"מי המתחרים המובילים בארה\"ב בתחום: {business_goal}? מה הם גובים?",
        f"מהי טווח המחירים המקובל בארה\"ב עבור: {business_goal}?",
        f"אילו פערים והזדמנויות עם השקעה נמוכה קיימים בתחום: {business_goal}?",
    ]


def researcher(client, question, idx, total):
    """חוקר: מחפש באינטרנט נתונים אמיתיים על שאלת מחקר אחת."""
    system = (
        "אתה חוקר שוק מקצועי. השתמש בחיפוש באינטרנט כדי למצוא נתונים אמיתיים "
        "ועדכניים על השוק האמריקאי: מחירים בדולרים, שמות מתחרים אמיתיים, מספרים, "
        "מגמות. צטט מקורות כשאפשר. אל תמציא מספרים — אם לא מצאת נתון, אמור זאת. "
        "כתוב את הממצאים בעברית, בצורה תמציתית ומסודרת עם נקודות."
    )
    return run_agent(client, system, question, use_web=True, max_tokens=8000)


def supplier(client, business_goal, findings):
    """
    ספק: מחפש ספקים ב-Alibaba/AliExpress ואתרי סיטונאות,
    ומחשב עלות נחיתה אמיתית (Landed Cost) לפריט.
    """
    context = "\n\n".join(f"ממצא {i+1}:\n{f}" for i, f in enumerate(findings))
    system = (
        "אתה מומחה רכש ושרשרת אספקה למכירה אונליין. "
        "חפש באינטרנט ב-Alibaba.com, AliExpress ואתרי סיטונאות נוספים. מצא: "
        "(1) ספקים ספציפיים + טווח מחירים לפריט (FOB price), "
        "(2) כמות מינימלית להזמנה (MOQ), "
        "(3) עלויות משלוח מסין לארה\"ב — אוויר (7-14 יום) מול ים (30-45 יום), "
        "(4) מיסי יבוא אמריקאיים (US import tariffs/duties) לקטגוריה הזו, "
        "(5) חישוב Landed Cost לפריט = מחיר מוצר + משלוח + מיסים. "
        "אם רלוונטי, כלול גם עלות Amazon FBA Fulfillment לפריט. "
        "השווה לפחות 2–3 אפשרויות ספקים שונות. "
        "כתוב הכל בעברית עם מספרים ספציפיים בדולרים ואחוזים."
    )
    user = (
        f"התחום/המוצר: {business_goal}\n\n"
        f"ממצאי מחקר שוק לקונטקסט:\n{context}\n\n"
        "חפש ספקים ומצא את אפשרות הנחיתה (Landed Cost) הכי אטרקטיבית."
    )
    return run_agent(client, system, user, use_web=True, max_tokens=10000)


def economist(client, business_goal, findings):
    """
    כלכלן: מחשב עלויות, תמחור ורווחיות על בסיס ממצאי המחקר.
    משתמש בחיפוש כדי למצוא עמלות פלטפורמות אמיתיות.
    """
    combined = "\n\n".join(f"ממצא {i+1}:\n{f}" for i, f in enumerate(findings))
    system = (
        "אתה כלכלן עסקי שמתמחה במכירה אונליין בארה\"ב. על בסיס ממצאי המחקר, "
        "בנה ניתוח רווחיות ריאליסטי בעברית. חפש באינטרנט עלויות אמיתיות: עמלות "
        "מרקטפלייסים (Amazon, Etsy, Shopify), עלויות משלוח, עלות פרסום ממוצעת "
        "(CPC/CPA), ועלות מוצר. הצג: (1) טבלת עלויות משוערת לפריט, (2) תמחור "
        "מומלץ, (3) רווח גולמי למכירה ובאחוזים, (4) כמה מכירות צריך כדי לכסות "
        "השקעה ראשונית נמוכה. ציין הנחות במפורש ואל תמציא מספרים שלא מצאת."
    )
    user = f"המטרה העסקית: {business_goal}\n\nממצאי המחקר:\n{combined}"
    return run_agent(client, system, user, use_web=True, max_tokens=10000)


def marketer(client, business_goal, findings):
    """
    שיווק: כותב תוכן שיווקי באנגלית לשוק האמריקאי, עם הסבר בעברית לכל חלק.
    זה הסוכן שמגשר על פער השפה — הלקוח מקבל אנגלית מוכנה להעתקה.
    """
    context = "\n\n".join(f"ממצא {i+1}:\n{f}" for i, f in enumerate(findings))
    system = (
        "אתה קופירייטר מומחה למסחר אלקטרוני בשוק האמריקאי. הלקוח אינו דובר "
        "אנגלית, לכן עליך לכתוב לו תוכן שיווקי מוכן לשימוש באנגלית, ולצדו הסבר "
        "ותרגום בעברית כדי שיבין מה כתבת. הפק עבור המוצר/התחום: "
        "(1) שם מוצר וכותרת מושכת באנגלית, "
        "(2) תיאור מוצר מלא באנגלית (לדף מכירה/אמזון), "
        "(3) 5 נקודות מכירה (bullet points) באנגלית, "
        "(4) טקסט קצר למודעת פרסום באנגלית, "
        "(5) רשימת מילות מפתח באנגלית (SEO). "
        "לכל חלק באנגלית — הוסף מיד אחריו תרגום/הסבר קצר בעברית בסוגריים. "
        "סדר את הכל עם כותרות ברורות."
    )
    user = (
        f"התחום/המוצר: {business_goal}\n\n"
        f"רקע מהמחקר (להתאמת המסרים לשוק):\n{context}"
    )
    return run_agent(client, system, user, use_web=False, max_tokens=12000)


def ceo_synthesize(client, business_goal, findings, supply, economics):
    """המנכ"ל מסכם הכל לדוח החלטות עם המלצה קונקרטית."""
    combined = "\n\n".join(f"### ממצאי מחקר {i+1}:\n{f}" for i, f in enumerate(findings))
    system = (
        "אתה מנכ\"ל שמכין דוח החלטות לבעלים של עסק אונליין שאינו דובר אנגלית. "
        "על בסיס המחקר, נתוני הספקים וניתוח הרווחיות, כתוב דוח ברור בעברית: "
        "(1) סיכום השוק והביקוש, (2) המתחרים והמחירים, (3) ניתוח שרשרת האספקה — "
        "מאיפה לקנות, מה עולה להגיע לארה\"ב, (4) ההזדמנות הכי מבטיחה עם השקעה "
        "נמוכה, (5) המלצה קונקרטית עם 3-5 צעדים ראשונים ריאליסטיים, "
        "(6) הערכת סיכונים ולוחות זמנים כנה. "
        "אם משהו לא משתלם — אמור זאת בכנות."
    )
    user = (
        f"המטרה העסקית: {business_goal}\n\n"
        f"ממצאי צוות המחקר:\n{combined}\n\n"
        f"נתוני ספקים ועלות נחיתה:\n{supply}\n\n"
        f"ניתוח רווחיות מהכלכלן:\n{economics}\n\n"
        "כתוב את דוח ההחלטות המלא."
    )
    return run_agent(client, system, user, use_web=False, max_tokens=16000)


# ---------------------------------------------------------------------------
# ייצוא ל-Word
# ---------------------------------------------------------------------------

def _add_md_content(doc, text):
    """ממיר markdown פשוט לפסקאות ב-Word."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], 3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], 2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], 1)
        elif stripped.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            set_rtl(p)
        elif stripped == "":
            p = doc.add_paragraph()
            set_rtl(p)
        else:
            clean = stripped.replace("**", "").replace("*", "")
            p = doc.add_paragraph(clean)
            set_rtl(p)


def create_word_report(business_goal, report, supply, economics, marketing, stamp):
    """יוצר קובץ Word מעוצב מהדוח. מחזיר שם הקובץ, או None אם הספרייה חסרה."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    doc = Document()

    # כותרת ראשית
    title = doc.add_heading(f"דוח עסקי: {business_goal}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub = doc.add_paragraph(f"נוצר: {stamp}  |  מודל: {MODEL}")
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    sections = [
        ("🧑‍💼 דוח החלטות (המנכ\"ל)", report),
        ("📦 ספקים ועלות נחיתה", supply),
        ("💰 ניתוח רווחיות (הכלכלן)", economics),
        ("✍️ תוכן שיווקי (סוכן השיווק)", marketing),
    ]

    for heading, content in sections:
        h = doc.add_heading(heading, 1)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_md_content(doc, content)
        doc.add_paragraph()

    docx_filename = f"report_{stamp}.docx"
    doc.save(docx_filename)
    return docx_filename


# ---------------------------------------------------------------------------
# פונקציה מרכזית — קוראת לכל הסוכנים בתור
# ---------------------------------------------------------------------------

def run_analysis(business_goal, on_progress=None):
    """
    מריץ את כל הסוכנים ומחזיר dict עם כל התוצאות + שמות קבצים.
    on_progress(step, msg) — callback אופציונלי לעדכוני התקדמות.
    """

    def _progress(step, msg):
        if on_progress:
            on_progress(step, msg)
        else:
            print(msg)

    client = get_client()

    _progress("plan", 'שלב 1 — המנכ"ל מפרק את המטרה למשימות מחקר...')
    tasks = ceo_plan(client, business_goal)

    findings = []
    for i, q in enumerate(tasks):
        _progress("research", f"שלב 2 — חוקר {i+1}/{len(tasks)}: {q}")
        findings.append(researcher(client, q, i + 1, len(tasks)))

    _progress("supply", "שלב 3 — סוכן הספקים מחפש מקורות ועלויות...")
    supply_data = supplier(client, business_goal, findings)

    _progress("econ", "שלב 4 — הכלכלן מחשב עלויות ורווחיות...")
    economics = economist(client, business_goal, findings)

    _progress("market", "שלב 5 — סוכן השיווק כותב תוכן שיווקי באנגלית...")
    marketing = marketer(client, business_goal, findings)

    _progress("report", 'שלב 6 — המנכ"ל מסכם את הדוח הסופי...')
    report = ceo_synthesize(client, business_goal, findings, supply_data, economics)

    stamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")
    md_filename = f"report_{stamp}.md"

    with open(md_filename, "w", encoding="utf-8") as f:
        f.write(f"# דוח עסקי מלא: {business_goal}\n\n")
        f.write(f"_נוצר ב-{stamp} | מודל: {MODEL}_\n\n")
        f.write("---\n\n")
        f.write("## 🧑‍💼 דוח החלטות (המנכ\"ל)\n\n" + report + "\n\n")
        f.write("## 📦 ספקים ועלות נחיתה\n\n" + supply_data + "\n\n")
        f.write("## 💰 ניתוח רווחיות (הכלכלן)\n\n" + economics + "\n\n")
        f.write("## ✍️ תוכן שיווקי מוכן (סוכן השיווק)\n\n" + marketing + "\n")

    _progress("export", "שומר קבצים...")
    docx_filename = create_word_report(
        business_goal, report, supply_data, economics, marketing, stamp
    )

    return {
        "report": report,
        "supply": supply_data,
        "economics": economics,
        "marketing": marketing,
        "md_filename": md_filename,
        "docx_filename": docx_filename,
        "stamp": stamp,
    }


# ---------------------------------------------------------------------------
# הרצה ראשית (שורת פקודה)
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("  מערכת סוכני AI לעסק אונליין בשוק האמריקאי")
    print(f"  (מודל בשימוש: {MODEL})")
    print("=" * 60)

    if len(sys.argv) > 1:
        business_goal = " ".join(sys.argv[1:])
    else:
        business_goal = input(
            "\nמה התחום/המטרה העסקית שתרצה לחקור?\n"
            "(לדוגמה: 'מכירת מוצרי טיפוח טבעיים אונליין בארה\"ב')\n> "
        ).strip()

    if not business_goal:
        print("לא הוזנה מטרה. יציאה.")
        return

    print(f"\nמטרה: {business_goal}\n")
    result = run_analysis(business_goal)

    print("\n" + "=" * 60)
    print(result["report"])
    print("=" * 60)
    print(f"\n[נשמר] דוח מלא:    {result['md_filename']}")
    if result["docx_filename"]:
        print(f"[נשמר] קובץ Word:   {result['docx_filename']}")
    print("פתח את הקבצים בכל עורך טקסט / Word כדי לראות הכל.")


if __name__ == "__main__":
    main()
